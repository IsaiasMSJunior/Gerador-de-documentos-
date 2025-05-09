import streamlit as st
import json
import firebase_admin
from firebase_admin import credentials, db
import pandas as pd
import calendar
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from datetime import datetime
from io import BytesIO
from openpyxl import load_workbook
from openpyxl.styles import PatternFill

# Adicione 'firebase-admin' em requirements.txt

# --- Inicialização do Firebase ---
service_account_info = json.loads(st.secrets["firebase_key"])
cred = credentials.Certificate(service_account_info)
firebase_admin.initialize_app(cred, {
    "databaseURL": st.secrets["databaseURL"]
})

# Configuração da página
st.set_page_config(page_title="Agenda Escolar", layout="wide")

# --- Helpers e Configurações Gerais ---
map_hor = {
    "1ª": "7:00–7:50", "2ª": "7:50–8:40", "3ª": "8:40–9:30",
    "4ª": "9:50–10:40", "5ª": "10:40–11:30", "6ª": "12:20–13:10", "7ª": "13:10–14:00"
}
meses = ["Janeiro","Fevereiro","Março","Abril","Maio","Junho",
         "Julho","Agosto","Setembro","Outubro","Novembro","Dezembro"]
ano_planej = 2025

# Funções para interação com Realtime Database

def get_db(path, default):
    data = db.reference(path).get()
    return data if data is not None else default


def set_db(path, value):
    db.reference(path).set(value)

# Funções auxiliares

def extrai_serie(turma: str) -> str:
    return turma[:-1]


def set_border(par: Paragraph):
    p = par._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bd = OxmlElement('w:bottom')
    bd.set(qn('w:val'),'single'); bd.set(qn('w:sz'),'4')
    bd.set(qn('w:space'),'1'); bd.set(qn('w:color'),'auto')
    pBdr.append(bd); pPr.append(pBdr)


def insert_after(par: Paragraph, text='') -> Paragraph:
    new_p = OxmlElement('w:p'); par._p.addnext(new_p)
    para = Paragraph(new_p, par._parent)
    if text: para.add_run(text)
    return para

# --- Funções de geração de documentos ---

def gerar_agenda_template(entries, df_bank, professor, semana, bimestre, cores_turmas):
    wb = load_workbook("agenda_modelo.xlsx")
    ws = wb.active
    ws["B1"] = professor
    ws["E1"] = semana
    row_map = {"1ª":4, "2ª":6, "3ª":8, "4ª":12, "5ª":14, "6ª":18, "7ª":20}
    col_map = {"Segunda":"C", "Terça":"D", "Quarta":"E", "Quinta":"F", "Sexta":"G"}
    for e in entries:
        col, row = col_map[e["dia"]], row_map[e["aula"]]
        ws[f"{col}{row}"] = f"{e['turma']} – {e['disciplina']}"
        fill = PatternFill(
            start_color=cores_turmas.get(e["turma"], "#FFFFFF").lstrip("#"),
            end_color=cores_turmas.get(e["turma"], "#FFFFFF").lstrip("#"),
            fill_type="solid"
        )
        ws[f"{col}{row}"].fill = fill
        titulo = ""
        if not df_bank.empty:
            sel = df_bank[
                (df_bank["DISCIPLINA"]==e["disciplina"]) &
                (df_bank["ANO/SÉRIE"]==extrai_serie(e["turma"])) &
                (df_bank["BIMESTRE"]==bimestre) &
                (df_bank["Nº da aula"]==e["num"]
            )]
            if not sel.empty:
                titulo = sel["TÍTULO DA AULA"].iloc[0]
        ws[f"{col}{row+1}"] = f"Aula {e['num']} – {titulo}"
        ws[f"{col}{row+1}"].fill = fill
    out = BytesIO(); wb.save(out); out.seek(0)
    return out


def gerar_plano_template(entries, df_bank, professor, semana, bimestre, turma,
                         metodologias, recursos, criterios, modelo="modelo_plano.docx"):
    doc = Document(modelo)
    header_disciplinas = ", ".join(sorted({e['disciplina'] for e in entries}))
    total_aulas = str(len(entries))
    # Cabeçalho
    for p in doc.paragraphs:
        for tag, value in {"ppp":professor,"ttt":turma,"sss":semana,"ddd":header_disciplinas,"nnn":total_aulas}.items():
            if tag in p.text:
                p.text = p.text.replace(tag, value)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for tag, value in {"ppp":professor,"ttt":turma,"sss":semana,"ddd":header_disciplinas,"nnn":total_aulas}.items():
                        if tag in p.text:
                            p.text = p.text.replace(tag, value)
    # Blocos de aula
    for p in doc.paragraphs:
        if p.text.strip() == "ccc":
            p.text = ""; last = p
            insert_after(last); set_border(last)
            for e in entries:
                sub = df_bank[
                    (df_bank["DISCIPLINA"]==e["disciplina"]) &
                    (df_bank["ANO/SÉRIE"]==extrai_serie(turma)) &
                    (df_bank["BIMESTRE"]==bimestre) &
                    (df_bank["Nº da aula"]==e["num"]
                )]
                titulo = sub["TÍTULO DA AULA"].iloc[0] if not sub.empty else ""
                pa = insert_after(last, f"Aula {e['num']} – {titulo}")
                pa.runs[0].bold=True; last=pa
                insert_after(last)
            # Extras: metodologias, recursos, criterios
            if metodologias:
                insert_after(last, "Metodologia:")
                for m in metodologias: insert_after(last, f"• {m}")
            if recursos:
                insert_after(last, "Recursos:")
                for r in recursos: insert_after(last, f"• {r}")
            if criterios:
                insert_after(last, "Critérios de Avaliação:")
                for c in criterios: insert_after(last, f"• {c}")
            break
    out = BytesIO(); doc.save(out); out.seek(0)
    return out


def gerar_guia_template(professor, turma, disciplina, bimestre, inicio, fim,
                        qtd_bimestre, qtd_semanal, metodologias, criterios,
                        df_bank, modelo="modelo_guia.docx"):
    doc = Document(modelo)
    replacements = {
        'ppp': professor,
        'ttt': turma,
        'bbb': bimestre,
        'iii': inicio.strftime('%d/%m/%Y'),
        'fff': fim.strftime('%d/%m/%Y'),
        'qqq': str(qtd_bimestre),
        'sss': str(qtd_semanal),
        'mmm': ", ".join(metodologias),
        'ccc': ", ".join(criterios),
        'ddd': disciplina
    }
    for p in doc.paragraphs:
        for tag, val in replacements.items():
            if tag in p.text: p.text = p.text.replace(tag, val)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for tag, val in replacements.items():
                        if tag in p.text: p.text = p.text.replace(tag, val)
    # Habilidades e objetos
    mask = (
        (df_bank["DISCIPLINA"]==disciplina)&
        (df_bank["ANO/SÉRIE"]==extrai_serie(turma))&
        (df_bank["BIMESTRE"]==bimestre)
    )
    habs = df_bank.loc[mask, "HABILIDADE"].dropna().astype(str).tolist()
    objs = df_bank.loc[mask, "OBJETO_DE_CONHECIMENTO"].dropna().astype(str).tolist()
    unique_habs = list(dict.fromkeys(habs))
    unique_objs = list(dict.fromkeys(objs))
    for p in doc.paragraphs:
        if 'hhh' in p.text: p.text = "\n".join(unique_habs)
        if 'ooo' in p.text: p.text = "\n".join(unique_objs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if 'hhh' in p.text: p.text = "\n".join(unique_habs)
                    if 'ooo' in p.text: p.text = "\n".join(unique_objs)
    out = BytesIO(); doc.save(out); out.seek(0)
    return out


def gerar_planejamento_template(professor, disciplina, turma, bimestre,
                                grupos, df_bank, modelo="modelo_planejamento.docx"):
    doc = Document(modelo)
    # Substituições iniciais
    for p in doc.paragraphs:
        for tag,val in {'ppp':professor,'ddd':disciplina,'ttt':turma,'bbb':bimestre}.items():
            if tag in p.text: p.text = p.text.replace(tag,val)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for tag,val in {'ppp':professor,'ddd':disciplina,'ttt':turma,'bbb':bimestre}.items():
                        if tag in p.text: p.text = p.text.replace(tag,val)
    # Grupos de planejamento
    for grp in grupos:
        p0 = doc.add_paragraph(f"Semana: {grp['semana']}")
        for n in grp['nums']:
            mask = (
                (df_bank["DISCIPLINA"]==disciplina)&
                (df_bank["ANO/SÉRIE"]==extrai_serie(turma))&
                (df_bank["BIMESTRE"]==bimestre)&
                (df_bank["Nº da aula"]==n)
            )
            titles = df_bank.loc[mask,"TÍTULO DA AULA"].dropna().tolist()
            title = titles[0] if titles else ""
            doc.add_paragraph(f"Aula {n} – {title}")
        doc.add_paragraph(f"Metodologia: {', '.join(grp['met'])}")
        doc.add_paragraph(f"Critérios: {', '.join(grp['crit'])}")
    out = BytesIO(); doc.save(out); doc.seek(0)
    return out

# --- Inicialização de estados via Firebase ---
if get_db("extras", None) is None:
    set_db("extras", {"metodologia":[],"recursos":[],"criterios":[]})
st.session_state.extras = get_db("extras", {"metodologia":[],"recursos":[],"criterios":[]})

if "professores" not in st.session_state:
    st.session_state.professores = get_db("professores", [])
if "turmas" not in st.session_state:
    st.session_state.turmas = get_db("turmas", {})
if "horarios" not in st.session_state:
    st.session_state.horarios = get_db("horarios", [])

# --- Sidebar ---
pages = [
    "Cadastro de Professor","Cadastro de Turmas","Cadastro de Horário",
    "Gerar Agenda e Plano","Cadastro Extras","Gerar Guia","Planejamento Bimestral"
]
for p in pages:
    if st.sidebar.button(p, use_container_width=True):
        st.session_state.page = p

# --- Páginas ---
if st.session_state.get('page') == "Cadastro de Professor":
    st.header("Cadastro de Professor")
    nome = st.text_input("Nome")
    disciplinas = st.multiselect("Disciplina(s)", [...])  # defina opções completas
    if st.button("Salvar Professor"):
        st.session_state.professores.append({"nome":nome,"disciplinas":disciplinas})
        set_db("professores", st.session_state.professores)
        st.success("Professor salvo!")
    st.write(st.session_state.professores)

elif st.session_state.get('page') == "Cadastro de Turmas":
    st.header("Cadastro de Turmas")
    # similar ao cadastro original, use set_db para salvar

elif st.session_state.get('page') == "Cadastro de Horário":
    st.header("Cadastro de Horário")
    # estrutura de horários, on button save: set_db("horarios", st.session_state.horarios)

elif st.session_state.get('page') == "Gerar Agenda e Plano":
    st.header("Gerar Agenda e Plano")
    # lógica igual, usa gerar_agenda_template e gerar_plano_template

elif st.session_state.get('page') == "Cadastro Extras":
    st.header("Cadastro Extras")
    # use set_db para atualizar extras

elif st.session_state.get('page') == "Gerar Guia":
    st.header("Gerar Guia")
    # lógica gerar_guia_template

elif st.session_state.get('page') == "Planejamento Bimestral":
    st.header("Planejamento Bimestral")
    # lógica gerar_planejamento_template
